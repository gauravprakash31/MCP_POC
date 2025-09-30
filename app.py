# ~/mcp-gdrive/app.py
import io, os, textwrap, streamlit as st
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document as DocxDocument

# ---------- CONFIG ----------
TOKEN = "/Users/sharvaripurighalla/mcp-gdrive/token.json"
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
MAX_BYTES = 120_000

@st.cache_resource
def _drive():
    c = Credentials.from_authorized_user_file(TOKEN, SCOPES)
    if not c.valid and c.refresh_token:
        c.refresh(Request())
    return build("drive", "v3", credentials=c)

def _download(req) -> bytes:
    buf = io.BytesIO()
    MediaIoBaseDownload(buf, req).next_chunk()
    return buf.getvalue()

def extract_text_for(meta: dict) -> str:
    """Universal extractor for Docs, Slides->PDF, Sheets->CSV, PDF, DOCX, TXT."""
    svc = _drive()
    fid, name, mime = meta["id"], meta["name"], meta["mimeType"]
    head = f"# {name} [{mime}]\n\n"

    if mime == "application/vnd.google-apps.document":
        data = _download(svc.files().export_media(fileId=fid, mimeType="text/plain"))
        body = data[:MAX_BYTES].decode("utf-8", "replace")

    elif mime == "application/vnd.google-apps.presentation":
        data = _download(svc.files().export_media(fileId=fid, mimeType="application/pdf"))
        text = pdf_extract_text(io.BytesIO(data))[:MAX_BYTES]
        body = text or "[No selectable text in slides export. (OCR is an optional addon.)]"

    elif mime == "application/vnd.google-apps.spreadsheet":
        data = _download(svc.files().export_media(fileId=fid, mimeType="text/csv"))
        csv_text = data[:MAX_BYTES].decode("utf-8", "replace")
        lines = csv_text.splitlines()
        preview = []
        for i, line in enumerate(lines):
            if i > 1000:
                preview.append("...[truncated additional rows]...")
                break
            preview.append(line.replace(",", "\t"))
        body = "\n".join(preview)

    elif mime == "application/pdf":
        data = _download(svc.files().get_media(fileId=fid))
        text = pdf_extract_text(io.BytesIO(data))[:MAX_BYTES]
        body = text or "[No selectable text; PDF may be scanned. (OCR optional addon.)]"

    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        data = _download(svc.files().get_media(fileId=fid))
        doc = DocxDocument(io.BytesIO(data))
        body = "\n".join(p.text for p in doc.paragraphs)[:MAX_BYTES]

    elif mime == "text/plain":
        data = _download(svc.files().get_media(fileId=fid))
        body = data[:MAX_BYTES].decode("utf-8", "replace")

    elif mime == "application/msword":
        try:
            data = _download(svc.files().export_media(
                fileId=fid,
                mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ))
            doc = DocxDocument(io.BytesIO(data))
            body = "\n".join(p.text for p in doc.paragraphs)[:MAX_BYTES]
        except Exception as e:
            body = f"[Legacy .doc conversion failed: {e}]"

    else:
        body = f"[Unsupported for direct text extraction: {name} ({mime}). Convert to Google Doc/PDF/TXT.]"

    return head + (body or "[No text found]")

def list_recent(n=10):
    svc = _drive()
    res = svc.files().list(pageSize=n, orderBy="modifiedTime desc",
                           q="trashed=false",
                           fields="files(id,name,mimeType,modifiedTime)").execute()
    return res.get("files", [])

def search_files(q, n=15):
    svc = _drive()
    res = svc.files().list(q=f"(name contains '{q}' or fullText contains '{q}') and trashed=false",
                           pageSize=n, fields="files(id,name,mimeType,modifiedTime)").execute()
    return res.get("files", [])

def summarize(text, title):
    # Prefer Anthropic (Claude), fallback OpenAI, else naive bullets
    ak = os.getenv("ANTHROPIC_API_KEY")
    if ak:
        from anthropic import Anthropic
        client = Anthropic(api_key=ak)
        msg = client.messages.create(
            model="claude-3-5-sonnet-latest",
            max_tokens=600, temperature=0.2,
            messages=[{"role": "user", "content":
                       f"Summarize in 5 bullets + 3 action items.\nTitle: {title}\nText:\n{text[:20000]}"}]
        )
        return msg.content[0].text.strip()
    ok = os.getenv("OPENAI_API_KEY")
    if ok:
        from openai import OpenAI
        o = OpenAI(api_key=ok)
        resp = o.chat.completions.create(
            model="gpt-4o-mini", temperature=0.2,
            messages=[{"role":"user","content":
                       f"Summarize in 5 bullets + 3 action items.\nTitle: {title}\nText:\n{text[:20000]}"}]
        )
        return resp.choices[0].message.content.strip()
    # naive fallback
    sents = [s.strip() for s in text.replace("\n"," ").split(". ") if s.strip()]
    bullets = "\n".join(f"- {textwrap.shorten(s, 140, placeholder='…')}" for s in sents[:5])
    actions = "\n".join(f"- {textwrap.shorten(s, 140, placeholder='…')}" for s in sents[5:8]) or "- (no clear actions)"
    return bullets + "\n\nActions:\n" + actions

# ---------- UI ----------
st.set_page_config(page_title="BOS • Drive Agent (POC)", layout="wide")
st.title("BOS • Google Drive AI Agent (POC)")

col = st.columns(3)
with col[0]:
    if st.button("List recent 10"):
        st.session_state["files"] = list_recent(10)
with col[1]:
    query = st.text_input("Search Drive", value=st.session_state.get("q",""))
with col[2]:
    if st.button("Search"):
        st.session_state["q"] = query
        st.session_state["files"] = search_files(query, 15)

files = st.session_state.get("files", [])
if not files:
    st.info("Click 'List recent 10' or run a search.")
else:
    for f in files:
        with st.expander(f"{f['name']} • {f['mimeType']} • {f.get('modifiedTime','')}"):
            c1, c2 = st.columns([1,1])
            with c1:
                if st.button("Extract text", key="t_"+f["id"]):
                    with st.spinner("Reading file…"):
                        text = extract_text_for(f)
                    st.session_state["text_"+f["id"]] = text
                text = st.session_state.get("text_"+f["id"])
                if text:
                    st.text_area("Content", text, height=280)
            with c2:
                text = st.session_state.get("text_"+f["id"])
                if text:
                    if st.button("Summarize", key="s_"+f["id"]):
                        with st.spinner("Summarizing…"):
                            st.session_state["sum_"+f["id"]] = summarize(text, f["name"])
                    if st.session_state.get("sum_"+f["id"]):
                        st.markdown("### Summary")
                        st.markdown(st.session_state["sum_"+f["id"]])
