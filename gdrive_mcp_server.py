# ~/mcp-gdrive/gdrive_mcp_server.py
import io, os, logging
from pathlib import Path
from mcp.server.fastmcp import FastMCP
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials

logging.basicConfig(level=logging.INFO)
mcp = FastMCP("gdrive")


TOKEN_PATH = Path("/Users/sharvaripurighalla/mcp-gdrive/token.json")
CREDS_PATH = Path("/Users/sharvaripurighalla/mcp-gdrive/credentials.json")
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]

def _creds():
    logging.info(f"[gdrive] using TOKEN_PATH={TOKEN_PATH}")
    if not TOKEN_PATH.exists():
        raise RuntimeError(f"Run `python auth_setup.py` first. Missing: {TOKEN_PATH}")
    c = Credentials.from_authorized_user_file(str(TOKEN_PATH), SCOPES)
    if not c.valid:
        if c.expired and c.refresh_token:
            c.refresh(Request())
            TOKEN_PATH.write_text(c.to_json())
        else:
            raise RuntimeError("Credentials invalid. Delete token.json and re-run auth_setup.py")
    return c

def _svc():
    return build("drive", "v3", credentials=_creds())

def _link(fid: str) -> str:
    return f"https://drive.google.com/open?id={fid}"

@mcp.tool()
def list_recent_files(limit: int = 5) -> str:
    """List recent files (name, id, link)."""
    res = _svc().files().list(
        pageSize=limit,
        orderBy="modifiedTime desc",
        q="trashed=false",
        fields="files(id,name,mimeType,modifiedTime)"
    ).execute()
    items = res.get("files", [])
    if not items:
        return "No files found."
    return "\n".join(f"- {f['name']} [id:{f['id']}] ({f['mimeType']}) → {_link(f['id'])}" for f in items)

@mcp.tool()
def search_files(query: str, limit: int = 5) -> str:
    """Search by name or contents."""
    q = f"(name contains '{query}' or fullText contains '{query}') and trashed=false"
    res = _svc().files().list(q=q, pageSize=limit, fields="files(id,name,mimeType)").execute()
    items = res.get("files", [])
    if not items:
        return f"No matches for {query!r}."
    return "\n".join(f"- {f['name']} [id:{f['id']}] ({f['mimeType']}) → {_link(f['id'])}" for f in items)

@mcp.tool()
def get_file_text(file_id: str, max_bytes: int = 40000) -> str:
    """Return text for Google Docs or .txt (truncated)."""
    s = _svc()
    meta = s.files().get(fileId=file_id, fields="id,name,mimeType").execute()
    mime, name = meta["mimeType"], meta["name"]
    buf = io.BytesIO()
    if mime == "application/vnd.google-apps.document":
        req = s.files().export_media(fileId=file_id, mimeType="text/plain")
    elif mime == "text/plain":
        req = s.files().get_media(fileId=file_id)
    else:
        return f"Unsupported for text extraction: {name} ({mime})."
    dl = MediaIoBaseDownload(buf, req)
    done = False
    while not done:
        _, done = dl.next_chunk()
    data = buf.getvalue()[:max_bytes]
    text = data.decode("utf-8", errors="replace")
    head = f"# {name}\n\n"
    if len(buf.getvalue()) > max_bytes:
        head += f"(Truncated to ~{max_bytes} bytes)\n\n"
    return head + text

@mcp.tool()
def debug_auth() -> str:
    """Show where creds are and if they exist."""
    return (
        f"CWD={os.getcwd()}\n"
        f"TOKEN_PATH={TOKEN_PATH}\n"
        f"TOKEN_EXISTS={TOKEN_PATH.exists()}\n"
        f"CRED_PATH={CREDS_PATH}\n"
        f"CRED_EXISTS={CREDS_PATH.exists()}"
    )
@mcp.tool()
def get_file_text_universal(file_id: str, max_bytes: int = 120_000) -> str:
    """
    Universal text extractor for Drive files.

    Supports:
      - Google Docs -> text/plain
      - Google Slides -> export PDF -> extract text (OCR optional)
      - Google Sheets -> export CSV -> return TSV-like preview
      - PDF -> extract text (OCR optional if scanned)
      - DOCX -> parse paragraphs
      - TXT -> direct download

    Returns a header + body. Truncates to ~max_bytes to stay chat-friendly.
    """
    import io
    from googleapiclient.http import MediaIoBaseDownload
    from pdfminer.high_level import extract_text as pdf_extract_text

    s = _svc()
    meta = s.files().get(fileId=file_id, fields="id,name,mimeType").execute()
    name, mime = meta["name"], meta["mimeType"]

    buf = io.BytesIO()

    def _download(req):
        # Simple one-chunk download 
        MediaIoBaseDownload(buf, req).next_chunk()
        return buf.getvalue()

    def _as_text(b: bytes) -> str:
        return b[:max_bytes].decode("utf-8", errors="replace")

    # ---- Google native types
    if mime == "application/vnd.google-apps.document":
        data = _download(s.files().export_media(fileId=file_id, mimeType="text/plain"))
        body = _as_text(data)

    elif mime == "application/vnd.google-apps.presentation":
        # Export to PDF and extract text
        data = _download(s.files().export_media(fileId=file_id, mimeType="application/pdf"))
        text = pdf_extract_text(io.BytesIO(data))[:max_bytes]
        if not text.strip():
            text = "[No selectable text found in exported Slides PDF. It may be image-based; enable OCR to read images.]"
        body = text

    elif mime == "application/vnd.google-apps.spreadsheet":
        # Export CSV (first sheet) and render as TSV-ish preview
        data = _download(s.files().export_media(fileId=file_id, mimeType="text/csv"))
        csv_text = _as_text(data)
        # Make it readable in chat: replace commas with tabs, cap rows
        lines = csv_text.splitlines()
        preview_rows = []
        for i, line in enumerate(lines):
            if i > 1000:
                preview_rows.append("...[truncated additional rows]...")
                break
            preview_rows.append(line.replace(",", "\t"))
        body = "\n".join(preview_rows)

    # ---- Binary/stored files
    elif mime == "application/pdf":
        data = _download(s.files().get_media(fileId=file_id))
        text = pdf_extract_text(io.BytesIO(data))[:max_bytes]
        if not text.strip():
            text = "[No selectable text extracted; PDF may be scanned. Enable OCR to read images.]"
        body = text

    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # DOCX
        data = _download(s.files().get_media(fileId=file_id))
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(data))
        body = "\n".join(p.text for p in doc.paragraphs)[:max_bytes]

    elif mime == "text/plain":
        data = _download(s.files().get_media(fileId=file_id))
        body = _as_text(data)

    elif mime == "application/msword":
        # Try exporting old .doc to DOCX then parse
        try:
            data = _download(s.files().export_media(
                fileId=file_id,
                mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ))
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(data))
            body = "\n".join(p.text for p in doc.paragraphs)[:max_bytes]
        except Exception as e:
            body = f"[Could not convert/read legacy .doc: {e}]"

    else:
        body = f"[Unsupported for direct text extraction: {name} ({mime}). Try converting to Google Doc/PDF/TXT.]"

    header = f"# {name} [{mime}]\n\n"
    return header + (body or "[No text found]")

def _ocr_pdf_to_text(pdf_bytes: bytes) -> str:
    # Lazy OCR pipeline: PDF -> images -> Tesseract
    from pdf2image import convert_from_bytes
    import pytesseract
    from PIL import Image
    pages = convert_from_bytes(pdf_bytes, dpi=200)
    parts = []
    for i, img in enumerate(pages[:10]):  # cap first 10 pages for cost/time
        parts.append(pytesseract.image_to_string(img))
    return "\n".join(parts)

if __name__ == "__main__":
    mcp.run(transport="stdio")


